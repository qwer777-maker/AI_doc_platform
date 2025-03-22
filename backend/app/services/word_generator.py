import os
from typing import List, Dict, Any, Optional
import logging
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from ..core.config import settings

from .ai_service_factory import AIServiceFactory

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class WordGenerator:
    """
    生成Word文档
    """
    def __init__(self, ai_service_type: str = "deepseek"):
        self.templates_dir = os.path.join(os.path.dirname(__file__), "../templates/word_templates")
        self.output_dir = os.path.abspath(os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "generated_docs"))
        logger.info(f"Word生成器输出目录: {self.output_dir}")
        os.makedirs(self.output_dir, exist_ok=True)
        
        # 确保目录权限正确
        try:
            os.chmod(self.output_dir, 0o777)
        except Exception as e:
            logger.warning(f"无法修改目录权限: {e}")
        
        self.ai_service = AIServiceFactory.create_service(ai_service_type)
    
    def generate(self, topic: str, outline: List[Dict[str, Any]], template_id: Optional[str] = None) -> Optional[str]:
        """
        根据大纲生成Word文档
        
        Args:
            topic: 文档主题
            outline: 文档大纲
            template_id: 可选的模板ID
            
        Returns:
            生成的Word文件路径，如果失败则返回None
        """
        try:
            logger.info(f"开始生成Word文档: 主题='{topic}', 章节数={len(outline)}")
            if template_id:
                logger.info(f"使用模板: {template_id}")
            
            # 创建文档
            doc = self._create_document(template_id)
            logger.info(f"创建新的Word文档" + (f" (使用模板: {template_id})" if template_id else ""))
            
            # 添加标题页
            logger.info("添加标题页")
            self._add_title_page(doc, topic)
            
            # 添加目录
            logger.info("添加目录")
            self._add_toc(doc)
            
            # 添加章节内容
            for section_index, section in enumerate(outline):
                section_title = section["title"]
                logger.info(f"处理章节 {section_index+1}/{len(outline)}: '{section_title}'")
                self._add_section(doc, section, topic)
                
                # 记录子章节信息
                subsections = section.get("subsections", [])
                if subsections:
                    logger.info(f"  章节 '{section_title}' 包含 {len(subsections)} 个子章节")
                    for i, subsection in enumerate(subsections):
                        subsection_title = subsection.get("title", "未知子章节")
                        logger.info(f"    子章节 {i+1}: '{subsection_title}'")
            
            # 添加参考文献
            logger.info("添加参考文献")
            self._add_references(doc, topic)
            
            # 保存文件
            file_name = f"{topic.replace(' ', '_')}_document.docx"
            file_path = os.path.join(self.output_dir, file_name)
            doc.save(file_path)
            
            # 计算页数（近似值）
            page_count = len(doc.paragraphs) // 20  # 假设每页约20个段落
            logger.info(f"Word文档生成成功: 约 {page_count} 页, 保存至: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"生成Word文档时出错: {str(e)}")
            return None
    
    def _create_document(self, template_id: Optional[str]) -> Document:
        """
        创建文档对象，可选择使用模板
        """
        # 如果提供了模板ID，尝试加载模板
        if template_id and template_id != "default":
            template_path = f"templates/word/{template_id}.docx"
            if os.path.exists(template_path):
                return Document(template_path)
        
        # 默认创建空白文档
        doc = Document()
        
        # 设置默认样式
        styles = doc.styles
        
        # 设置正文样式
        style_normal = styles['Normal']
        font = style_normal.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # 设置标题样式
        for i in range(1, 4):
            style_name = f'Heading {i}'
            if style_name in styles:
                style = styles[style_name]
                font = style.font
                font.name = 'Arial'
                font.size = Pt(16 - (i-1)*2)
                font.bold = True
        
        return doc
    
    def _add_title_page(self, doc: Document, topic: str) -> None:
        """
        添加标题页
        """
        # 添加标题
        doc.add_paragraph().add_run().add_break()  # 添加空行
        title = doc.add_paragraph(topic)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        
        # 添加副标题
        doc.add_paragraph().add_run().add_break()  # 添加空行
        subtitle = doc.add_paragraph("AI自动生成的文档")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.italic = True
        
        # 添加日期
        doc.add_paragraph().add_run().add_break()  # 添加空行
        date = doc.add_paragraph()
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        import datetime
        date.add_run(datetime.datetime.now().strftime("%Y年%m月%d日"))
        
        # 添加分页符
        doc.add_page_break()
    
    def _add_toc(self, doc: Document) -> None:
        """
        添加目录
        """
        doc.add_paragraph("目录", style='Heading 1')
        
        # 添加目录域代码
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        
        # 使用域代码添加目录
        # 这只是一个占位符，实际的目录需要在Word中更新
        run.add_text("目录将在Word中显示。请右键点击并选择'更新域'来显示目录。")
        
        # 添加分页符
        doc.add_page_break()
    
    def _add_section(self, doc: Document, section: Dict[str, Any], topic: str) -> None:
        """
        添加章节
        """
        # 添加章节标题
        section_title = section["title"]
        doc.add_paragraph(section_title, style='Heading 1')
        
        # 使用AI服务生成章节内容
        content = self.ai_service.generate_section_content(topic, section_title, "word")
        
        # 添加章节介绍
        intro = doc.add_paragraph()
        intro.add_run(content)
        
        # 添加子章节
        subsections = section.get("subsections", [])
        if not subsections:
            # 如果没有子章节，添加一些默认内容
            self._add_default_content(doc, section_title, topic)
            return
        
        for subsection in subsections:
            self._add_subsection(doc, subsection, section_title, topic)
    
    def _add_subsection(self, doc: Document, subsection: Dict[str, Any], section_title: str, topic: str) -> None:
        """
        添加子章节
        """
        # 添加子章节标题
        subsection_title = subsection["title"]
        doc.add_paragraph(subsection_title, style='Heading 2')
        
        # 使用AI服务生成子章节内容
        content = self.ai_service.generate_section_content(topic, subsection_title, "word")
        
        # 将内容分段添加
        paragraphs = content.strip().split('\n\n')
        for p_text in paragraphs:
            if p_text:
                p = doc.add_paragraph()
                p.add_run(p_text.strip())
                p.paragraph_format.line_spacing = 1.5
    
    def _add_default_content(self, doc: Document, section_title: str, topic: str) -> None:
        """
        添加默认内容
        """
        content = f"""
{section_title}是{topic}的核心组成部分，它包含多个关键要素和特性。

深入理解{section_title}对掌握整个主题至关重要。通过系统分析其基本原理、应用场景和发展趋势，我们可以更全面地把握{topic}的本质。

{section_title}的理论基础建立在多年的研究和实践之上。众多学者和专家通过实证研究和理论探索，不断丰富和完善相关知识体系。

在实际应用中，{section_title}展现出强大的适应性和实用价值。无论是在教育、商业还是技术创新领域，都能找到其成功应用的案例。

未来研究将进一步探索{section_title}的潜力和应用前景。随着新技术和新方法的出现，我们有理由相信，{section_title}将在{topic}的发展中发挥更加重要的作用。
        """
        
        paragraphs = content.strip().split('\n\n')
        for p_text in paragraphs:
            if p_text:
                p = doc.add_paragraph()
                p.add_run(p_text.strip())
                p.paragraph_format.line_spacing = 1.5
    
    def _add_references(self, doc: Document, topic: str) -> None:
        """
        添加参考文献
        """
        doc.add_page_break()
        doc.add_paragraph("参考文献", style='Heading 1')
        
        # 添加一些模拟的参考文献
        references = [
            f"Smith, J. (2022). 理解{topic}的基本原理. 学术期刊, 45(2), 112-128.",
            f"Johnson, A., & Williams, B. (2021). {topic}的应用与实践. 科技出版社.",
            f"Chen, L., Wang, H., & Zhang, Y. (2023). {topic}的最新进展. 研究评论, 10(3), 78-95.",
            f"Taylor, M. (2020). {topic}在教育领域的应用. 教育研究, 33(1), 45-62.",
            f"Brown, R., & Davis, S. (2022). {topic}的未来发展趋势. 未来研究, 15(4), 201-215."
        ]
        
        for ref in references:
            p = doc.add_paragraph()
            p.add_run(ref)
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.line_spacing = 1.5 